"""
reference_loader.py
-------------------
Loads and caches reference documents from data/references/ for injection
into LLM prompts.

Currently supported:
  - Verification Methods.docx  → plain text, injected into Stage 2 & 3 prompts

Usage:
    from src.utils.reference_loader import load_verification_methods
    vm_context = load_verification_methods()   # cached after first call
"""

from __future__ import annotations

import logging
from functools import lru_cache
from pathlib import Path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Default path — resolved relative to the repo root
# ---------------------------------------------------------------------------

_REPO_ROOT = Path(__file__).resolve().parents[2]
_DEFAULT_VM_PATH = _REPO_ROOT / "data" / "references" / "Verification Methods.docx"
_DEFAULT_RT_PATH = _REPO_ROOT / "data" / "references" / "requirement_type_context.md"

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_docx_text(path: Path) -> str:
    """
    Extract all paragraph text from a .docx file, preserving table cell text.
    Returns a plain-text string with paragraphs separated by newlines.
    """
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to read .docx reference files. "
            "Install it with:  pip install python-docx"
        ) from exc

    doc = Document(str(path))
    lines: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables (each cell on its own line, blank line between rows)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                lines.append(" | ".join(row_cells))
        lines.append("")  # blank separator between tables

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@lru_cache(maxsize=1)
def load_verification_methods(path: str | Path | None = None) -> str:
    """
    Load and return the full text of Verification Methods.docx.

    Parameters
    ----------
    path : optional override for the .docx file location.
           Defaults to data/references/Verification Methods.docx
           (resolved relative to the repo root).

    Returns
    -------
    Plain-text string of the document contents, cached after the first call.
    Returns an empty string and logs a warning if the file cannot be read.
    """
    target = Path(path) if path else _DEFAULT_VM_PATH

    if not target.exists():
        logger.warning(
            "Verification Methods reference not found at '%s'. "
            "Prompts will be sent without verification method context.",
            target,
        )
        return ""

    try:
        text = _extract_docx_text(target)
        logger.info("Loaded verification methods reference (%d chars) from '%s'.", len(text), target)
        return text
    except Exception as exc:
        logger.warning("Failed to read '%s': %s — proceeding without context.", target, exc)
        return ""


def verification_methods_block(path: str | Path | None = None) -> str:
    """
    Return the verification methods text wrapped in a clearly delimited block
    suitable for direct injection into a prompt string.

    Returns an empty string if the reference file is unavailable.
    """
    text = load_verification_methods(path)
    if not text:
        return ""
    return (
        "VERIFICATION METHODS REFERENCE (from Verification Methods.docx)\n"
        "Use the definitions below when selecting a verification_method.\n"
        "════════════════════════════════════════════════════════════════\n"
        f"{text}\n"
        "════════════════════════════════════════════════════════════════\n"
    )

@lru_cache(maxsize=1)
def load_requirement_type_context(path: str | Path | None = None) -> str:
    target = Path(path) if path else _DEFAULT_RT_PATH
    if not target.exists():
        logger.warning("Requirement type context not found at '%s'.", target)
        return ""
    try:
        text = target.read_text(encoding="utf-8")
        logger.info("Loaded requirement type context (%d chars) from '%s'.", len(text), target)
        return text
    except Exception as exc:
        logger.warning("Failed to read '%s': %s — proceeding without context.", target, exc)
        return ""

def requirement_type_block(path: str | Path | None = None) -> str:
    text = load_requirement_type_context(path)
    if not text:
        return ""
    return (
        "REQUIREMENT TYPE REFERENCE\n"
        "Use the definitions below when selecting a requirement_type.\n"
        "════════════════════════════════════════════════════════════════\n"
        f"{text}\n"
        "════════════════════════════════════════════════════════════════\n"
    )