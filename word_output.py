import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from pathlib import Path
import pandas as pd
from docx.oxml import OxmlElement

def add_multilevel_numbering(doc):
    numbering = doc.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "1")

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")

        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")

        lvlText = OxmlElement("w:lvlText")
        # ✅ Correct hierarchical numbering
        lvlText.set(
            qn("w:val"),
            ".".join(f"%{i+1}" for i in range(level + 1)) + "."
        )

        pStyle = OxmlElement("w:pStyle")
        # ✅ CORRECT Word style IDs (with space)
        pStyle.set(qn("w:val"), f"Heading {level + 1}")

        lvl.extend([start, numFmt, lvlText, pStyle])
        abstract.append(lvl)

    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")

    abstractId = OxmlElement("w:abstractNumId")
    abstractId.set(qn("w:val"), "1")

    num.append(abstractId)
    numbering.append(num)

# ─────────────────────────────────────────────
# Helpers for font enforcement (CRITICAL)
# ─────────────────────────────────────────────

def force_style_font(style, name="Arial", size=None, bold=None):
    font = style.font
    font.name = name
    if size is not None:
        font.size = size
    if bold is not None:
        font.bold = bold

    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()

    for attr in (
        "w:asciiTheme",
        "w:hAnsiTheme",
        "w:eastAsiaTheme",
        "w:cstheme",
    ):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]

    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def add_paragraph_arial(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


# ─────────────────────────────────────────────
# Text processing helpers (UNCHANGED LOGIC)
# ─────────────────────────────────────────────

def _ensure_lines(cleaned_or_text):
    if isinstance(cleaned_or_text, str):
        lines = cleaned_or_text.splitlines()
    else:
        lines = list(cleaned_or_text)

    norm = []
    for ln in lines:
        ln = str(ln).strip()
        ln = re.sub(r"^/\*+|\*+/$", "", ln).strip()
        ln = re.sub(r"^\*+\s?", "", ln).strip()
        norm.append(ln)
    return norm


def extract_module_basename_from_cleaned(cleaned_or_text) -> str:
    lines = _ensure_lines(cleaned_or_text)
    for ln in lines:
        m = re.match(r"^\s*Filename\s*:\s*<?([^>\r\n]+)>?\s*$", ln, flags=re.IGNORECASE)
        if m:
            filename = m.group(1).strip()
            base = re.sub(r"\.c$", "", filename, flags=re.IGNORECASE).strip()
            return base or "Module"
    return "Module"


def extract_section_text_from_cleaned(cleaned_or_text, section_name: str) -> str:
    lines = _ensure_lines(cleaned_or_text)

    start = None
    for i, ln in enumerate(lines):
        if re.match(rf"^\s*{re.escape(section_name)}\s*:\s*$", ln, flags=re.IGNORECASE):
            start = i + 1
            break
    if start is None:
        return ""

    while start < len(lines) and re.match(r"^\s*-{3,}\s*$", lines[start]):
        start += 1

    out = []
    for ln in lines[start:]:
        if re.match(r"^\s*[A-Za-z][A-Za-z0-9 _\(\)]*\s*:\s*$", ln):
            break
        if re.match(r"^\s*-{3,}\s*$", ln):
            continue
        if ln.strip():
            out.append(re.sub(r"\s+", " ", ln.strip()))
    return "\n".join(out).strip()


def extract_all_functions_from_cleaned(cleaned_or_text):
    lines = _ensure_lines(cleaned_or_text)

    start = None
    for i, ln in enumerate(lines):
        if re.match(r"^\s*Function\(s\)\s*:\s*$", ln, flags=re.IGNORECASE) or \
           re.match(r"^\s*Functions\s*:\s*$", ln, flags=re.IGNORECASE):
            start = i + 1
            break
    if start is None:
        return []

    while start < len(lines) and re.match(r"^\s*-{3,}\s*$", lines[start]):
        start += 1

    items = []
    for ln in lines[start:]:
        if re.match(r"^\s*[A-Za-z][A-Za-z0-9 _\(\)]*\s*:\s*$", ln):
            break
        if re.match(r"^\s*-{3,}\s*$", ln):
            continue

        ln = ln.strip()
        if not ln:
            continue
        ln = re.sub(r"^\-\s*", "", ln)
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln:
            items.append(ln)
    return items

def remove_paragraph_spacing(style):
    pPr = style.element.get_or_add_pPr()

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")      # single spacing
    spacing.set(qn("w:lineRule"), "auto")

# ─────────────────────────────────────────────
# MAIN DOCUMENT GENERATION
# ─────────────────────────────────────────────

def generate_doc_from_cleaned_and_df(df, cleaned, out_path: str, debug: bool = False):
    doc = Document()
    add_multilevel_numbering(doc)
    remove_paragraph_spacing(doc.styles["List Bullet"])
    remove_paragraph_spacing(doc.styles["Normal"])

    # Normal + lists
    force_style_font(doc.styles["Normal"], "Arial", size=Pt(11))
    force_style_font(doc.styles["List Bullet"], "Arial", size=Pt(11))

    # Headings
    HEADING_SIZES = {1: 14, 2: 12, 3: 10}
    for level, size in HEADING_SIZES.items():
        style = doc.styles[f"Heading {level}"]
        force_style_font(style, "Arial", size=Pt(size), bold=True)
        style.font.color.rgb = RGBColor(0, 0, 0)

    # ── Extract content ──
    lines = _ensure_lines(cleaned)
    module = extract_module_basename_from_cleaned(lines)
    desc = extract_section_text_from_cleaned(lines, "Description")
    hdr_funcs = extract_all_functions_from_cleaned(lines)

    # ── Document content ──
    doc.add_heading("Capability Requirements", level=1)
    doc.add_heading(f"{module} Module Capability Requirements", level=2)

    add_paragraph_arial(doc, desc if desc else "(No description found)")

    add_paragraph_arial(
        doc,
        f"The P3 MCP {module} Module shall provide the following capabilities:"
    )

    if hdr_funcs:
        for _, row in df.iterrows():
            fn = str(row["Function"]).strip()
            p = doc.add_paragraph(fn, style="List Bullet")
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
    else:
        add_paragraph_arial(doc, "(No functions found in header)")

    # ── Function requirements ──
    for _, row in df.iterrows():
        fn = str(row["Function"]).strip()
        doc.add_heading(fn, level=2)

        req = row.get("Requirements", "")
        if isinstance(req, list):
            req_list = [str(x).strip() for x in req if str(x).strip()]
        else:
            req_list = [r.strip() for r in str(req).splitlines() if r.strip()]

        if req_list:
            for r in req_list:
                add_paragraph_arial(doc, r)
        else:
            add_paragraph_arial(doc, "(No requirements)")

    doc.save(out_path)